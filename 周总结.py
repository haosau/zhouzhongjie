from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def generate_weekly_summary():
    # 读取模板文档
    template_path = "C:/备份/学习计划和总结/黄浩20230626-20230702 周学习总结与下周计划.docx"
    doc = Document(template_path)

    # 获取当前日期，计算当前周的周一到周日日期
    today = datetime.today()
    start_of_week = today - timedelta(days=today.weekday())
    dates = [start_of_week + timedelta(days=i) for i in range(7)]

    # 格式化日期为字符串
    date_strings = [date.strftime("%Y.%m.%d") for date in dates]
    date_range = f"{date_strings[0]}——{date_strings[-1]}"
    weekdays = ['周一', '周二', '周三', '周四', '周五', '周六', '周日']
    date_weekday_strings = [f"{date_strings[i]} {weekdays[i]}" for i in range(7)]

    # 更新文档标题和表格中的日期
    # 更新文档标题和表格中的日期
    for paragraph in doc.paragraphs:
        if "周学习总结与下周计划" in paragraph.text:
            paragraph.text = f"{date_range} 周学习总结与下周计划"
        elif "周工作总结：" in paragraph.text:
            paragraph.text = f"{date_range} 周工作总结："

    for table in doc.tables:
        for i, row in enumerate(table.rows[1:]):  # 跳过标题行
            cell = row.cells[0]
            cell.text = date_weekday_strings[i]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 生成新的文件名
    new_filename = f"黄浩{date_strings[0].replace('.', '')}-{date_strings[-1].replace('.', '')} 周学习总结与下周计划.docx"
    output_dir = "C:/备份/学习计划和总结"
    new_filepath = os.path.join(output_dir, new_filename)

    # 保存新的文档
    doc.save(new_filepath)

    return new_filepath


# 运行函数并生成文档
new_file_path = generate_weekly_summary()
print(f"New document created: {new_file_path}")
